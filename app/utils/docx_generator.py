"""Generate a professional resume DOCX from plain text using python-docx."""
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_HEADER_RE = re.compile(
    r"^\s*(SUMMARY|OBJECTIVE|EXPERIENCE|WORK EXPERIENCE|EDUCATION|SKILLS|"
    r"CERTIFICATIONS?|PROJECTS?|ACHIEVEMENTS?|AWARDS?|PUBLICATIONS?|"
    r"LANGUAGES?|INTERESTS?|CONTACT(?: INFO(?:RMATION)?)?|REFERENCES?|VOLUNTEER|HOBBIES?)\s*$",
    re.IGNORECASE,
)


def _is_section_header(line: str) -> bool:
    return bool(_HEADER_RE.match(line)) or (line.isupper() and 3 < len(line.strip()) < 60)


def _add_bottom_border(para, color_hex: str = "1a1a2e", sz_eighths: int = 8):
    """Add a bottom border to a paragraph using OpenXML."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz_eighths))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _para(doc: Document, text: str, bold=False, size_pt=10.0,
          color="#222222", space_before=0.0, space_after=2.0,
          italic=False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.color.rgb = _rgb(color)
    return p


def generate_resume_docx(resume_text: str) -> bytes:
    """Convert plain-text resume to a professionally formatted DOCX and return bytes."""
    doc = Document()

    # Reset default Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    # Page margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    lines = resume_text.splitlines()
    first_content = True
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # ── Name (first non-empty line) ─────────────────────────────────
        if first_content:
            p = _para(doc, line.strip(), bold=True, size_pt=20, color="#1a1a2e", space_after=3)
            first_content = False
            i += 1

            # Collect contact info lines
            contact_parts = []
            while i < len(lines):
                cline = lines[i].rstrip()
                if not cline.strip() or _is_section_header(cline):
                    break
                contact_parts.append(cline.strip())
                i += 1

            if contact_parts:
                _para(doc, " | ".join(contact_parts), size_pt=9, color="#555555", space_after=4)

            # Thick rule under header block
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(2)
            sep.paragraph_format.space_after = Pt(6)
            _add_bottom_border(sep, color_hex="#1a1a2e", sz_eighths=12)
            continue

        # ── Section header ──────────────────────────────────────────────
        if _is_section_header(line):
            # Small gap before section
            gap = doc.add_paragraph()
            gap.paragraph_format.space_before = Pt(6)
            gap.paragraph_format.space_after = Pt(0)

            hp = doc.add_paragraph()
            hp.paragraph_format.space_before = Pt(0)
            hp.paragraph_format.space_after = Pt(4)
            run = hp.add_run(line.strip().upper())
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = _rgb("#1a1a2e")
            _add_bottom_border(hp, color_hex="#aaaaaa", sz_eighths=4)
            i += 1
            continue

        # ── Bullet points ───────────────────────────────────────────────
        if line.lstrip().startswith(("•", "-", "*", "–", "·")):
            clean = re.sub(r"^[\s•\-\*–·]+", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.12)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"•  {clean.strip()}")
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = _rgb("#222222")
            i += 1
            continue

        # ── Regular body text ───────────────────────────────────────────
        _para(doc, line.strip(), size_pt=10, color="#333333", space_after=2)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
