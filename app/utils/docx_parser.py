"""Extract plain text from a DOCX file using python-docx."""
import io
from docx import Document


def extract_text(file_path) -> str:
    """Extract all text from a DOCX file, preserving paragraph structure."""
    with open(file_path, "rb") as f:
        doc = Document(f)

    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    row_parts.append(cell_text)
            if row_parts:
                lines.append(" | ".join(row_parts))

    return "\n".join(lines)
