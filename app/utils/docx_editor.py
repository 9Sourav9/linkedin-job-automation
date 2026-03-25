"""
Edit an existing DOCX resume by replacing text while preserving all formatting.

Strategy (line-by-line):
  1. Split original + tailored into paired lines.
  2. For each changed line, find the matching paragraph in the DOCX.
  3. Replace ONLY the text content of that paragraph, keeping every run's
     font name, font size, bold, italic, colour, alignment, and spacing
     exactly as-is.

python-docx represents a paragraph as a list of "runs", each with their own
character-level formatting. We replace text by putting the new content into
the first run and clearing the rest — so the visual style of the first run
is retained, which is typically the same for the whole line.
"""
import io
import logging
import re
import unicodedata

from docx import Document

logger = logging.getLogger(__name__)


def _norm(text: str) -> str:
    """Normalize unicode + collapse whitespace for comparison."""
    text = unicodedata.normalize("NFKD", text)
    return re.sub(r"\s+", " ", text).strip()


def _replace_para_text(para, new_text: str) -> None:
    """
    Replace all text in a paragraph with new_text.
    The first run's character formatting (font, size, bold, italic, color)
    is preserved. All other runs are cleared.
    """
    if not para.runs:
        # No runs — add one
        para.add_run(new_text)
        return

    # Save the first run's full XML-level formatting via its rPr element
    # python-docx keeps per-run formatting in run._r (the XML element)
    # Easiest: just set text on first run, blank out the rest
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _get_all_paras(doc: Document):
    """Yield all paragraphs in the document including those inside tables."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def apply_section_changes_docx(docx_bytes: bytes, changes: list[dict]) -> bytes:
    """
    Apply section-level changes to a DOCX file, preserving all layout and formatting.

    For each change, we match each original line to a paragraph in the document
    and replace just the text content.

    Raises RuntimeError if no replacements were made (so caller can fall back).
    """
    doc = Document(io.BytesIO(docx_bytes))

    # Build map: normalised_original_line -> new_line
    line_map: dict[str, str] = {}
    for change in changes:
        orig_lines = [l.strip() for l in change.get("original", "").splitlines() if l.strip()]
        new_lines  = [l.strip() for l in change.get("tailored", "").splitlines() if l.strip()]
        for orig, new in zip(orig_lines, new_lines):
            if _norm(orig) != _norm(new):
                line_map[_norm(orig)] = new

    if not line_map:
        logger.info("docx_editor: no differing lines — returning original unchanged")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    replaced = 0

    for para in _get_all_paras(doc):
        para_text = para.text.strip()
        if not para_text:
            continue

        para_norm = _norm(para_text)

        # Strategy 1: exact normalised match
        if para_norm in line_map:
            logger.debug("docx_editor exact match: '%s'", para_text[:60])
            _replace_para_text(para, line_map[para_norm])
            replaced += 1
            continue

        # Strategy 2: fuzzy word-overlap match (≥ 75% word overlap)
        para_words = set(para_norm.lower().split())
        if len(para_words) < 3:
            continue
        for orig_norm, new_text in line_map.items():
            orig_words = set(orig_norm.lower().split())
            if len(orig_words) < 3:
                continue
            overlap = len(para_words & orig_words) / max(len(orig_words), 1)
            if overlap >= 0.75:
                logger.debug("docx_editor fuzzy (%.0f%%): '%s'", overlap * 100, para_text[:60])
                _replace_para_text(para, new_text)
                replaced += 1
                break

    logger.info("docx_editor: %d/%d lines replaced", replaced, len(line_map))

    if replaced == 0:
        raise RuntimeError("docx_editor: could not find any matching paragraphs to replace")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
