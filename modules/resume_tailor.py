"""
Resume tailoring using Claude API (claude-opus-4-6).
Reads base resume DOCX, tailors it to a specific JD, saves new DOCX.
"""
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import anthropic
import config


client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)


def _read_docx_text(path: str) -> str:
    """Extract all text from a DOCX file preserving structure."""
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    # Also grab table content
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return "\n".join(lines)


def _docx_to_structured(path: str) -> list[dict]:
    """Return paragraphs with style info for reconstruction."""
    doc = Document(path)
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append({
            "text": para.text,
            "style": para.style.name,
            "alignment": para.alignment,
            "runs": [
                {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_size": run.font.size.pt if run.font.size else None,
                    "font_color": (
                        str(run.font.color.rgb)
                        if run.font.color and run.font.color.type
                        else None
                    ),
                }
                for run in para.runs
            ],
        })
    return paragraphs


def tailor_resume(
    job_title: str,
    company: str,
    job_description: str,
    job_id: str,
) -> str:
    """
    Use Claude to tailor the base resume for a specific job.
    Returns path to tailored resume DOCX.
    """
    os.makedirs(config.RESUMES_DIR, exist_ok=True)

    base_resume_text = _read_docx_text(config.BASE_RESUME_PATH)
    base_paragraphs = _docx_to_structured(config.BASE_RESUME_PATH)

    prompt = f"""You are an expert resume writer. Your task is to tailor the candidate's base resume for a specific job posting.

## Candidate's Base Resume:
{base_resume_text}

## Target Job:
**Title:** {job_title}
**Company:** {company}

**Job Description:**
{job_description}

## Instructions:
1. Analyze the job description carefully — identify required skills, keywords, tools, and qualifications.
2. Rewrite and optimize the resume to highlight relevant experience and skills that match the JD.
3. Use keywords from the JD naturally throughout the resume (for ATS optimization).
4. Keep the same overall structure and factual content (don't invent experience).
5. Strengthen bullet points to show measurable impact relevant to this role.
6. Adjust the professional summary/objective to target this specific role and company.
7. Reorder skills/technologies to prioritize what's most relevant to this JD.
8. Keep the resume to 1-2 pages worth of content.

## Output Format:
Return ONLY the tailored resume content, preserving the same section structure as the original.
Use the exact same section headers as the original resume.
Each section should be clearly separated.
Do not add any explanation or preamble — just the resume text.
"""

    print(f"  [Claude] Tailoring resume for: {job_title} at {company}...")

    with client.messages.stream(
        model="claude-opus-4-6",
        max_tokens=4096,
        thinking={"type": "adaptive"},
        messages=[{"role": "user", "content": prompt}],
    ) as stream:
        tailored_text = ""
        for block in stream:
            pass
        final = stream.get_final_message()
        for block in final.content:
            if block.type == "text":
                tailored_text = block.text
                break

    # Build the tailored DOCX
    output_path = _build_tailored_docx(tailored_text, base_paragraphs, job_title, company, job_id)
    print(f"  [Resume] Saved tailored resume: {output_path}")
    return output_path


def _build_tailored_docx(
    tailored_text: str,
    base_paragraphs: list[dict],
    job_title: str,
    company: str,
    job_id: str,
) -> str:
    """Build a DOCX from tailored text, preserving formatting from base resume."""
    # Load original doc as template to preserve formatting/styles
    doc = Document(config.BASE_RESUME_PATH)

    # Clear existing paragraphs
    for para in doc.paragraphs:
        para.clear()

    # Parse tailored text into sections
    lines = tailored_text.strip().split("\n")

    # Map base paragraph styles by position for reuse
    base_styles = {p["style"] for p in base_paragraphs}

    # Clear all content and rebuild
    # Remove all paragraphs from body
    body = doc.element.body
    for child in list(body):
        body.remove(child)

    # Add content
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Detect section headers (ALL CAPS or starts with common headers)
        is_header = (
            stripped.isupper()
            or stripped.upper() == stripped
            or re.match(
                r"^(SUMMARY|EXPERIENCE|EDUCATION|SKILLS|PROJECTS|CERTIFICATIONS|OBJECTIVE|PROFILE|WORK HISTORY)",
                stripped.upper(),
            )
        )

        if is_header and len(stripped) < 60:
            p = doc.add_paragraph(stripped)
            p.style = "Heading 2" if "Heading 2" in base_styles else "Normal"
            run = p.runs[0] if p.runs else p.add_run(stripped)
            run.bold = True
        elif stripped.startswith(("•", "-", "*", "·")):
            # Bullet point
            p = doc.add_paragraph(stripped.lstrip("•-*· "), style="List Bullet")
        elif re.match(r"^\d{4}", stripped) or "|" in stripped:
            # Date or divider line → preserve as-is
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.bold = True
        else:
            doc.add_paragraph(stripped)

    # Save
    safe_company = re.sub(r"[^\w\-_]", "_", company)[:30]
    safe_title = re.sub(r"[^\w\-_]", "_", job_title)[:30]
    filename = f"resume_{job_id}_{safe_title}_{safe_company}.docx"
    output_path = os.path.join(config.RESUMES_DIR, filename)
    doc.save(output_path)
    return output_path


def get_resume_bytes(resume_path: str) -> bytes:
    """Read resume file as bytes for upload."""
    with open(resume_path, "rb") as f:
        return f.read()
